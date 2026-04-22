import PyPDF2
from docx import Document
import re
import os
from datetime import datetime

class TextExtractor:
    """Extrait le texte des fichiers PDF et DOCX"""
    
    @staticmethod
    def extract_pdf(file_path):
        """Extrait le texte d'un fichier PDF"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text()
            return text.lower()
        except Exception as e:
            raise Exception(f"Erreur extraction PDF: {str(e)}")
    
    @staticmethod
    def extract_docx(file_path):
        """Extrait le texte d'un fichier DOCX"""
        try:
            doc = Document(file_path)
            text = ""
            for para in doc.paragraphs:
                text += para.text + " "
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
            return text.lower()
        except Exception as e:
            raise Exception(f"Erreur extraction DOCX: {str(e)}")
    
    @staticmethod
    def extract_text(file_path):
        """Extrait le texte en fonction du type de fichier"""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return TextExtractor.extract_pdf(file_path)
        elif file_ext == '.docx':
            return TextExtractor.extract_docx(file_path)
        else:
            raise Exception(f"Format de fichier non supporté: {file_ext}")
    
    @staticmethod
    def clean_text(text):
        """Nettoie le texte pour l'analyse"""
        # Supprime les caractères spéciaux inutiles
        text = re.sub(r'[^\w\s\-àâäéèêëïîôùûüœæç]', ' ', text)
        # Supprime les espaces multiples
        text = re.sub(r'\s+', ' ', text)
        return text.strip()
    
    @staticmethod
    def extract_sections(text):
        """Extrait les sections principales du CV"""
        sections = {
            'formation': '',
            'experience': '',
            'competences': '',
            'langues': '',
            'full_text': text
        }
        
        # Regex pour détecter les sections
        patterns = {
            'formation': r'(formation|diplôme|études|cursus|scolarité)(.*?)(?=\n\n|[a-z]+:|$)',
            'experience': r'(expérience|parcours professionnel|carrière|emplois)(.*?)(?=\n\n|[a-z]+:|$)',
            'competences': r'(compétence|savoir-faire|skills|aptitudes)(.*?)(?=\n\n|[a-z]+:|$)',
            'langues': r'(langue|langues|language)(.*?)(?=\n\n|[a-z]+:|$)'
        }
        
        for section, pattern in patterns.items():
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                sections[section] = match.group(0)
        
        return sections
    
    @staticmethod
    def extract_dates(text):
        """Extrait les périodes de travail du texte
        
        Reconnaît les formats :
        - 2018 - 2024
        - 2018-2024
        - janvier 2018 - mars 2024
        - Jan 2018 / Mar 2024
        - 01/2018 - 03/2024
        """
        
        # Mapper les mois en français/anglais vers numéro
        months_map = {
            'janvier': 1, 'jan': 1, 'february': 2, 'février': 2, 'fev': 2,
            'mars': 3, 'mar': 3, 'april': 4, 'avril': 4, 'avr': 4,
            'may': 5, 'mai': 5, 'june': 6, 'juin': 6, 'juin': 6,
            'july': 7, 'juillet': 7, 'juil': 7, 'august': 8, 'août': 8,
            'september': 9, 'septembre': 9, 'sept': 9, 'october': 10, 'octobre': 10,
            'oct': 10, 'november': 11, 'novembre': 11, 'nov': 11, 'december': 12,
            'décembre': 12, 'dec': 12
        }
        
        periods = []
        
        # Pattern 1: Année-Année (2018-2024 ou 2018 - 2024)
        pattern1 = r'(\d{4})\s*[-–/]\s*(\d{4})'
        for match in re.finditer(pattern1, text, re.IGNORECASE):
            start_year = int(match.group(1))
            end_year = int(match.group(2))
            if 1950 <= start_year <= datetime.now().year and start_year <= end_year <= datetime.now().year:
                periods.append((start_year, 1, end_year, 12))
        
        # Pattern 2: Mois Année - Mois Année (janvier 2018 - mars 2024)
        pattern2 = r'(\w+)\s+(\d{4})\s*[-–/]\s*(?:à|to|au)?\s*(\w+)\s+(\d{4})'
        for match in re.finditer(pattern2, text, re.IGNORECASE):
            start_month_str = match.group(1).lower()
            start_year = int(match.group(2))
            end_month_str = match.group(3).lower()
            end_year = int(match.group(4))
            
            start_month = months_map.get(start_month_str, 1)
            end_month = months_map.get(end_month_str, 12)
            
            if 1950 <= start_year <= datetime.now().year and start_year <= end_year <= datetime.now().year:
                periods.append((start_year, start_month, end_year, end_month))
        
        # Pattern 3: MM/YYYY - MM/YYYY (01/2018 - 03/2024)
        pattern3 = r'(\d{1,2})/(\d{4})\s*[-–/]\s*(\d{1,2})/(\d{4})'
        for match in re.finditer(pattern3, text, re.IGNORECASE):
            start_month = int(match.group(1))
            start_year = int(match.group(2))
            end_month = int(match.group(3))
            end_year = int(match.group(4))
            
            if 1 <= start_month <= 12 and 1 <= end_month <= 12:
                if 1950 <= start_year <= datetime.now().year and start_year <= end_year <= datetime.now().year:
                    periods.append((start_year, start_month, end_year, end_month))
        
        return periods
    
    @staticmethod
    def calculate_total_experience_years(periods):
        """Calcule la durée totale d'expérience en années
        
        Args:
            periods: Liste de tuples (start_year, start_month, end_year, end_month)
        
        Returns:
            float: Nombre total d'années d'expérience
        """
        if not periods:
            return 0
        
        total_months = 0
        
        for start_year, start_month, end_year, end_month in periods:
            # Convertir en mois depuis une date de référence
            start_date = datetime(start_year, start_month, 1)
            # Pour la fin, prendre le dernier jour du mois
            if end_month == 12:
                end_date = datetime(end_year + 1, 1, 1)
            else:
                end_date = datetime(end_year, end_month + 1, 1)
            
            delta = end_date - start_date
            months = delta.days / 30.44  # Moyenne de jours par mois
            total_months += months
        
        # Convertir en années
        total_years = total_months / 12
        
        return round(total_years, 1)
